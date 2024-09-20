import psycopg2
from docx import Document

conn = psycopg2.connect(
    dbname="TESTS",
    user="postgres",
    password="postgres",
    host="localhost",
    port="5432"
)
cursor = conn.cursor()

# Execute a consulta para obter todas as linhas da tabela
cursor.execute("SELECT * FROM memoriais")
rows = cursor.fetchall()

# Caminho para o arquivo modelo
modelo_path = "E:\POSTGRE_MEMORIAL/Modelo_REURB1.docx"

# Para cada linha, abra o arquivo modelo e acrescente os dados
for i, row in enumerate(rows):
    doc = Document(modelo_path)  # Abre o arquivo modelo
    try:
        style = doc.styles['Normal']
        doc.add_heading(f"Linha {i}", level=1, style=style)
    except KeyError:
        doc.add_heading(f"Linha {i}", level=1)
    for j, value in enumerate(row):
        doc.add_paragraph(f"Coluna {j}: {value}")
    doc.save(f"E:\\POSTGRE_MEMORIAL\\memoriais\\linha_{i}.docx")
# Feche a conexão
cursor.close()
conn.close()

# # Conecte-se ao banco de dados PostgreSQL
# conn = psycopg2.connect(
#     dbname="ALPHAVIDA",
#     user="postgres",
#     password="postgres",
#     host="localhost",
#     port="5432"
# )
# cursor = conn.cursor()

# # Execute a consulta para obter todas as linhas da tabela
# cursor.execute("SELECT * FROM memoriais")
# rows = cursor.fetchall()

# # Caminho para o arquivo modelo
# modelopath = "E:\POSTGRE_MEMORIAL\Modelo_memo_REURB.docx"

# # Para cada linha, abra o arquivo modelo e acrescente os dados
# for i, row in enumerate(rows):
#     doc = Document(modelopath)  # Abre o arquivo modelo
#     # doc.add_heading(f"Linha {i}", level=1)
#     style = doc.styles._get_style_id_from_style('Heading 1', None)
#     if style:
#         doc.add_heading(f"Linha {i}", level=1, style=style)
#     else:
#         doc.add_heading(f"Linha {i}", level=1)
#     for j, value in enumerate(row):
#         doc.add_paragraph(f"Coluna {j}: {value}")
#     doc.save(f"E:\POSTGRE_MEMORIAL\memoriais\linha_{i}.docx")

# # Feche a conexão
# cursor.close()
# conn.close()


# # Conecte-se ao banco de dados PostgreSQL
# conn = psycopg2.connect(
#     dbname="TESTS",
#     user="postgres",
#     password="postgres",
#     host="localhost",
#     port="5432"
# )
# cursor = conn.cursor()

# # Execute a consulta para obter todas as linhas da tabela
# cursor.execute("SELECT * FROM memoriais")
# rows = cursor.fetchall()

# # Para cada linha, crie um arquivo DOCX
# for i, row in enumerate(rows):
#     doc = Document()
#     doc.add_heading(f"Linha {i}", level=1)
#     for j, value in enumerate(row):
#         doc.add_paragraph(f"Coluna {j}: {value}")
#     doc.save(f"E:\POSTGRE_MEMORIAL\memoriais\linha_{i}.docx")

# # Feche a conexão
# cursor.close()
# conn.close()
